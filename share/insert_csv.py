import os

from oneLabProject import settings

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "oneLabProject.settings")

import django
django.setup()

import pandas as pd
from django.db import transaction
from file.models import File
from share.models import Share, ShareFile, ShareFileContent
import random
from docx import Document
from openpyxl import Workbook

# CSV 파일을 DataFrame으로 로드
csv_file_path = './datasets/total_data.csv'
df = pd.read_csv(csv_file_path)

# 파일 삽입
@transaction.atomic
def insert_files():
    for index, row in df.iterrows():
        # Share 모델에 데이터 추가
        grades = ['1학년', '2학년', '3학년', '4학년']
        grade = random.choice(grades)

        majors = ['경영/경제', '컴퓨터공학', '전기/전자', '건축/토목', '재료/산업', '농림/축산', '사회/과학', '간호/보건',
                  '의예/의학', '화공/에너지', '화학/생명/환경', '수학/물리', '교육', '언어/문학', '인문학', '미술/예술', '생활과학', '음악/영화', '체육/무용']
        major = random.choice(majors)

        university_members = [9, 11, 12, 13, 14, 15, 16]
        university_member = random.choice(university_members)
        share_instance = Share.objects.create(
            share_title=f'제목 테스트 {index}',
            share_content=f'내용 테스트 {index}',
            share_points=1000,
            share_choice_major=major,
            share_choice_grade=grade,
            share_type='과제',
            share_text_major='경제학과',
            share_text_name=f'2학기 기말고사 {index}',
            share_post_status='1',
            university_id=university_member,
        )

        # 파일 정보 삽입
        content = row['Content']
        file_name = f'file_{index}'  # 파일 이름 생성
        extensions = ['.hwp', '.docx', '.xlsx']
        extension = random.choice(extensions)
        file_path = f'../upload/share/2024/05/23/{file_name}{extension}'  # 파일이 저장될 경로
        if extension == '.hwp':
            with open(file_path, 'wb') as f:
                f.write(content.encode())  # 파일 내용을 바이너리로 변환하여 저장
        elif extension == '.docx':
            # docx 파일 형태로 저장
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
        else:
            # xlsx 파일 형태로 저장
            wb = Workbook()
            ws = wb.active
            ws.append([content])
            wb.save(file_path)

        # File 모델에 파일 정보 삽입
        file_size = os.path.getsize(file_path)  # 파일 크기 계산
        file_instance = File.objects.create(
            file_size=file_size,
        )

        # ShareFile 모델에 파일 정보 삽입
        share_file_instance = ShareFile.objects.create(
            file=file_instance,
            path=file_path,
            name=file_name,
            share=share_instance,
            # file_extension=extension,
        )

        def get_file_content(file_path, file_id):
            # ImageFieldFile 객체를 문자열로 변환
            file_path_str = str(file_path)

            # 파일 경로에서 "../upload/" 제거
            if file_path_str.startswith("../upload/"):
                file_path_str = file_path_str.replace("../upload/", "")

            # 절대 경로 생성
            file_path_full = os.path.join(settings.MEDIA_ROOT, file_path_str)

            if file_path_str.lower().endswith('.hwp'):
                with open(file_path_full, 'r', encoding='utf-8') as file:
                    file_content = file.read()
                return file_content
            elif file_path_str.lower().endswith('.docx'):
                file_content = get_docx_content(file_path)
                return file_content
            elif file_path_str.lower().endswith('.xlsx'):
                file_content = get_excel_content(file_path)
                return file_content
            else:
                return '다름'

        def get_excel_content(file_path):
            try:
                # 엑셀 파일을 읽어서 DataFrame으로 변환
                df = pd.read_excel(file_path)
                # DataFrame을 문자열로 변환하여 반환
                file_content = df.to_string()
                return file_content
            except FileNotFoundError:
                print(f"File not found: {file_path}")
                return None
            except Exception as e:
                print(f"Error in reading file: {file_path}")
                print(e)
                return None

        def get_docx_content(file_path):
            try:
                doc = Document(file_path)
                # 문단별로 텍스트를 추출하여 리스트로 저장
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                # 리스트를 하나의 문자열로 결합하여 반환
                file_content = '\n'.join(paragraphs)
                return file_content
            except FileNotFoundError:
                print(f"File not found: {file_path}")
                return None
            except Exception as e:
                print(f"Error in reading file: {file_path}")
                print(e)
                return None

        # ShareFileContent 모델에 파일 정보 삽입
        share_file_text = ShareFileContent.objects.create(
            share=share_instance,
            text= get_file_content(share_file_instance.path, share_file_instance),
            file_name=share_file_instance.name,
        )

# 파일 삽입 함수 실행
insert_files()
